import os
import json
from datetime import datetime
from werkzeug.utils import secure_filename
from flask import current_app

from models import db
from models.document import Document


def get_file_info(file_path, file_type):
    info = {}
    file_size = os.path.getsize(file_path)

    if file_type == 'pdf':
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            info['page_count'] = len(reader.pages)
        except Exception:
            info['page_count'] = 0

    elif file_type == 'docx':
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            info['paragraph_count'] = len(doc.paragraphs)
        except Exception:
            info['paragraph_count'] = 0

    elif file_type == 'pptx':
        try:
            from pptx import Presentation
            prs = Presentation(file_path)
            info['slide_count'] = len(prs.slides)
        except Exception:
            info['slide_count'] = 0

    elif file_type == 'xlsx':
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True)
            sheets = []
            total_rows = 0
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                row_count = ws.max_row or 0
                sheets.append({'name': sheet_name, 'rows': row_count})
                total_rows += row_count
            info['sheets'] = sheets
            info['total_rows'] = total_rows
        except Exception:
            info['sheets'] = []
            info['total_rows'] = 0

    elif file_type == 'txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                info['line_count'] = content.count('\n') + 1
                info['char_count'] = len(content)
        except Exception:
            info['line_count'] = 0
            info['char_count'] = 0

    info['file_size'] = file_size
    return info


def allowed_file(filename):
    if not filename or '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext in current_app.config['ALLOWED_EXTENSIONS']


def save_uploaded_file(file, asset_id):
    if file and allowed_file(file.filename):
        original_filename = file.filename
        ext = original_filename.rsplit('.', 1)[1].lower()
        timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S%f')
        filename = f"{asset_id}_{timestamp}_{secure_filename(original_filename)}"
        upload_folder = current_app.config['UPLOAD_FOLDER']
        os.makedirs(upload_folder, exist_ok=True)
        file_path = os.path.join(upload_folder, filename)
        file.save(file_path)

        file_info = get_file_info(file_path, ext)
        info_json = json.dumps(file_info, default=str)

        doc = Document(
            asset_id=asset_id,
            filename=filename,
            original_filename=original_filename,
            file_type=ext,
            file_size=file_info.get('file_size', 0),
            file_info=info_json
        )
        db.session.add(doc)
        db.session.commit()
        return doc
    return None
